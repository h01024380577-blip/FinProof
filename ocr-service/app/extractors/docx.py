"""
DOCX extraction via python-docx.

Unlike FinProof's legacy TS path (JSZip + regex XML stripping, which loses
table structure), this preserves table cells by serializing each row as TSV.
A successfully parsed DOCX is high-confidence (0.97) because there is no OCR
uncertainty involved.
"""

from __future__ import annotations

import io

DOCX_CONFIDENCE = 0.97


def extract(body: bytes) -> dict:
    try:
        from docx import Document
    except Exception as error:  # pragma: no cover - dependency guard
        return {
            "text": "",
            "confidence": 0.0,
            "provider": "python-docx",
            "pages": 0,
            "has_tables": False,
            "warnings": [f"python-docx unavailable: {error}"],
        }

    try:
        document = Document(io.BytesIO(body))
    except Exception as error:
        return {
            "text": "",
            "confidence": 0.0,
            "provider": "python-docx",
            "pages": 0,
            "has_tables": False,
            "warnings": [f"docx open failed: {error}"],
        }

    parts: list[str] = []

    for paragraph in document.paragraphs:
        if paragraph.text and paragraph.text.strip():
            parts.append(paragraph.text.strip())

    has_tables = False
    for table in document.tables:
        has_tables = True
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append("\t".join(cells))

    text = "\n".join(parts).strip()
    warnings: list[str] = []
    confidence = DOCX_CONFIDENCE
    if not text:
        confidence = 0.0
        warnings.append("no text extracted from docx")

    return {
        "text": text,
        "confidence": confidence,
        "provider": "python-docx",
        "pages": 0,
        "has_tables": has_tables,
        "warnings": warnings,
    }
