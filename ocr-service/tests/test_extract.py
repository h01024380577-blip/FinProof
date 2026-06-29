"""
Smoke tests for the OCR service.

Run inside a venv (see README):
    pip install -r requirements.txt pytest httpx
    pytest

Tesseract-dependent cases are skipped automatically when the binary is absent,
so digital-PDF and DOCX coverage still runs without a system Tesseract install.
"""

from __future__ import annotations

import io
import shutil

import pytest
from fastapi.testclient import TestClient

from app.main import app

client = TestClient(app)

TESSERACT_AVAILABLE = shutil.which("tesseract") is not None


def _digital_pdf(text: str) -> bytes:
    import fitz

    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), text)
    data = doc.tobytes()
    doc.close()
    return data


def _docx_with_table() -> bytes:
    from docx import Document

    document = Document()
    document.add_paragraph("연 10% 확정 수익 보장")
    table = document.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "구분"
    table.rows[0].cells[1].text = "금리"
    table.rows[1].cells[0].text = "기본"
    table.rows[1].cells[1].text = "3.0%"
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _low_res_image() -> bytes:
    from PIL import Image, ImageDraw

    image = Image.new("RGB", (220, 60), "white")
    draw = ImageDraw.Draw(image)
    draw.text((5, 20), "FINPROOF TEST", fill="black")
    buffer = io.BytesIO()
    image.save(buffer, format="PNG")
    return buffer.getvalue()


def test_health() -> None:
    response = client.get("/health")
    assert response.status_code == 200
    assert response.json() == {"status": "ok"}


def test_digital_pdf_extraction() -> None:
    pdf = _digital_pdf("FinProof digital pdf body text")
    response = client.post(
        "/extract",
        files={"file": ("sample.pdf", pdf, "application/pdf")},
    )
    assert response.status_code == 200
    body = response.json()
    assert body["text"].strip() != ""
    assert "FinProof" in body["text"]
    assert 0.0 <= body["confidence"] <= 1.0
    assert body["confidence"] >= 0.82  # digital text layer is high-confidence
    assert body["provider"] in {"pymupdf", "pdfplumber"}


def test_docx_preserves_table_cells() -> None:
    docx = _docx_with_table()
    response = client.post(
        "/extract",
        files={
            "file": (
                "sample.docx",
                docx,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert response.status_code == 200
    body = response.json()
    assert body["provider"] == "python-docx"
    assert body["has_tables"] is True
    assert "금리" in body["text"]  # table cell content survives
    assert body["confidence"] >= 0.82


@pytest.mark.skipif(not TESSERACT_AVAILABLE, reason="tesseract binary not installed")
def test_image_ocr_returns_scaled_confidence() -> None:
    image = _low_res_image()
    response = client.post(
        "/extract",
        files={"file": ("scan.png", image, "image/png")},
    )
    assert response.status_code == 200
    body = response.json()
    assert body["provider"] == "tesseract"
    assert 0.0 <= body["confidence"] <= 1.0


def test_corrupted_file_is_low_confidence_not_500() -> None:
    response = client.post(
        "/extract",
        files={"file": ("broken.pdf", b"%PDF-1.4 not really a pdf", "application/pdf")},
    )
    assert response.status_code == 200  # never a 500
    body = response.json()
    assert body["confidence"] == 0.0
    assert body["text"] == ""
    assert body["warnings"]


def test_unsupported_type_falls_back_to_low_confidence() -> None:
    response = client.post(
        "/extract",
        files={"file": ("data.bin", b"\x00\x01\x02binary", "application/octet-stream")},
    )
    assert response.status_code == 200
    body = response.json()
    assert body["confidence"] == 0.0
